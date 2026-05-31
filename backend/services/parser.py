# backend/services/parser.py

# fitz: PyMuPDF's Python interface for reading PDF files
# PyMuPDF is one of the fastest PDF libraries available in Python
import fitz  # PyMuPDF

# python-docx: reads Microsoft Word .docx files
import docx

# re: Python's built-in regular expression module
# Used for text cleaning and contact info extraction
import re

# io: creates in-memory file-like objects from bytes
# Allows us to parse file bytes without saving to disk
import io

# logging: standard Python logging (replaces print in production)
import logging

# typing: type hints for function signatures
from typing import Tuple, Optional

# Our Pydantic models for validated data structures
from backend.models.resume import ResumeCreate, ResumeMetadata

# Configure module-level logger
# __name__ = "backend.services.parser" — shows in log output
logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Production resume parser supporting PDF and DOCX formats.
    
    Design decisions:
    - Stateless: no instance variables, all methods take explicit inputs
    - Defensive: every operation wrapped in try/except with informative errors
    - In-memory: parses bytes directly, never writes temp files to disk
    - Extractive: preserves both raw and cleaned text for auditability
    """

    # Allowed MIME types — used for validation before parsing attempt
    ALLOWED_EXTENSIONS = {".pdf", ".docx"}

    # Maximum file size: 10MB in bytes (10 * 1024 * 1024)
    MAX_FILE_SIZE = 10 * 1024 * 1024

    def parse(self, file_bytes: bytes, filename: str) -> ResumeCreate:
        """
        Main entry point. Dispatches to PDF or DOCX parser based on extension.
        
        Args:
            file_bytes: Raw binary content of the uploaded file
            filename: Original filename including extension (e.g. "resume_john.pdf")
            
        Returns:
            ResumeCreate: Fully populated Pydantic model ready for MongoDB
            
        Raises:
            ValueError: If file type unsupported or file is corrupt/empty
        """

        # Validate file size before doing any processing
        # Fail fast: don't waste CPU parsing a 500MB file
        if len(file_bytes) > self.MAX_FILE_SIZE:
            raise ValueError(
                f"File size {len(file_bytes)} bytes exceeds maximum "
                f"{self.MAX_FILE_SIZE} bytes ({self.MAX_FILE_SIZE // (1024*1024)}MB)"
            )

        # Extract file extension and normalize to lowercase
        # Example: "Resume_John.PDF" → ".pdf"
        # os.path.splitext("resume.pdf") returns ("resume", ".pdf")
        import os
        _, extension = os.path.splitext(filename.lower())

        # Validate file type against whitelist
        if extension not in self.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{extension}'. "
                f"Allowed types: {self.ALLOWED_EXTENSIONS}"
            )

        # Dispatch to appropriate parser based on file type
        if extension == ".pdf":
            raw_text, page_count = self._parse_pdf(file_bytes)
        elif extension == ".docx":
            raw_text, page_count = self._parse_docx(file_bytes)

        # Validate extracted text — catch corrupt or image-only PDFs
        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError(
                "Could not extract meaningful text from the file. "
                "The file may be image-based, password-protected, or corrupt. "
                "Please ensure the resume contains selectable text."
            )

        # Clean the raw text for NLP processing
        cleaned_text = self._clean_text(raw_text)

        # Extract contact information from cleaned text
        # Returns None if not found — never crashes on missing data
        name = self._extract_name(cleaned_text)
        email = self._extract_email(cleaned_text)
        phone = self._extract_phone(cleaned_text)

        # Count words in cleaned text for metadata
        word_count = len(cleaned_text.split())

        # Build the metadata object
        metadata = ResumeMetadata(
            filename=filename,
            file_type=extension.lstrip("."),  # ".pdf" → "pdf"
            file_size_bytes=len(file_bytes),
            page_count=page_count,
            word_count=word_count
        )

        # Assemble and return the complete resume model
        # Pydantic validates all fields here — if anything is wrong, it raises
        return ResumeCreate(
            candidate_name=name,
            candidate_email=email,
            candidate_phone=phone,
            raw_text=raw_text,
            cleaned_text=cleaned_text,
            metadata=metadata
        )

    def _parse_pdf(self, file_bytes: bytes) -> Tuple[str, int]:
        """
        Extracts text from a PDF file using PyMuPDF.
        
        PyMuPDF reads the PDF's content stream directly and reconstructs
        text in reading order, handling multi-column layouts better than
        most alternatives.
        
        Args:
            file_bytes: Raw PDF file bytes
            
        Returns:
            Tuple of (extracted_text, page_count)
        """
        try:
            # Open PDF from bytes in memory — no temp file on disk
            # fitz.open() with stream parameter reads from bytes
            # filetype="pdf" tells PyMuPDF what format to expect
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")

            # List to accumulate text from all pages
            text_pages = []

            # Iterate over every page (zero-indexed)
            for page_number in range(len(pdf_document)):
                # Load the page object
                page = pdf_document[page_number]

                # Extract text from this page
                # sort=True: sorts text blocks by reading order (top-left to bottom-right)
                # This is critical for resume parsing where order carries meaning
                page_text = page.get_text(sort=True)

                # Only add non-empty pages
                if page_text.strip():
                    text_pages.append(page_text)

            # Record page count before closing
            page_count = len(pdf_document)

            # Always close the PDF document to free memory
            pdf_document.close()

            # Join all pages with double newline to preserve section boundaries
            full_text = "\n\n".join(text_pages)

            logger.info(f"PDF parsed: {page_count} pages, {len(full_text)} characters")

            return full_text, page_count

        except Exception as e:
            # Re-raise with more context for upstream error handling
            raise ValueError(f"PDF parsing failed: {str(e)}")

    def _parse_docx(self, file_bytes: bytes) -> Tuple[str, int]:
        """
        Extracts text from a DOCX file using python-docx.
        
        A .docx file is a ZIP containing XML. python-docx parses the XML
        and provides a Python object model over paragraphs and tables.
        
        Args:
            file_bytes: Raw DOCX file bytes
            
        Returns:
            Tuple of (extracted_text, section_count)
        """
        try:
            # Wrap bytes in a BytesIO object — python-docx expects a file-like object
            # BytesIO makes bytes look like a file without disk I/O
            docx_stream = io.BytesIO(file_bytes)

            # Open the document from the stream
            document = docx.Document(docx_stream)

            # List to accumulate all text content
            text_parts = []

            # Extract text from paragraphs
            # Paragraphs include headings, bullet points, and regular text
            for paragraph in document.paragraphs:
                # paragraph.text joins all "runs" (text spans) within the paragraph
                # A run is a contiguous sequence of characters with the same formatting
                para_text = paragraph.text.strip()

                # Skip empty paragraphs (common in Word docs)
                if para_text:
                    text_parts.append(para_text)

            # Extract text from tables
            # Resumes often use tables for skills sections or two-column layouts
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)

            # Count sections as a proxy for page count in DOCX
            section_count = len(document.sections)

            # Join all parts with newlines
            full_text = "\n".join(text_parts)

            logger.info(
                f"DOCX parsed: {len(document.paragraphs)} paragraphs, "
                f"{len(full_text)} characters"
            )

            return full_text, section_count

        except Exception as e:
            raise ValueError(f"DOCX parsing failed: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """
        Normalizes raw extracted text for NLP processing.
        
        Cleaning pipeline (order matters):
        1. Unicode normalization
        2. Remove non-printable characters
        3. Normalize whitespace
        4. Remove excessive punctuation
        5. Normalize line breaks
        
        Args:
            text: Raw text from PDF or DOCX parser
            
        Returns:
            Cleaned, normalized text string
        """
        # Step 1: Normalize Unicode characters
        # unicodedata.normalize converts accented chars to ASCII equivalents
        # 'NFKD' decomposes characters (é → e + combining accent)
        # encode('ascii', 'ignore') drops non-ASCII bytes
        # decode('ascii') converts bytes back to string
        import unicodedata
        text = unicodedata.normalize('NFKD', text)
        text = text.encode('ascii', 'ignore').decode('ascii')

        # Step 2: Remove non-printable control characters
        # \x00-\x08: null and other control chars
        # \x0b-\x0c: vertical tab, form feed
        # \x0e-\x1f: shift-out through unit separator
        # \x7f: delete character
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)

        # Step 3: Replace multiple spaces with single space
        # \s+ matches one or more whitespace characters
        # But preserve newlines — they denote section boundaries
        text = re.sub(r'[ \t]+', ' ', text)

        # Step 4: Replace 3+ consecutive newlines with exactly 2
        # This preserves section separation without excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Step 5: Remove lines that are just punctuation or symbols
        # These are often decorative separators (===, ---, ...)
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            # Keep line if it has at least 2 alphabetic characters
            # re.search(r'[a-zA-Z]{2,}', ...) checks for real words
            if stripped and re.search(r'[a-zA-Z]{2,}', stripped):
                cleaned_lines.append(stripped)

        # Step 6: Rejoin cleaned lines
        text = '\n'.join(cleaned_lines)

        # Step 7: Final strip of leading/trailing whitespace
        return text.strip()

    def _extract_email(self, text: str) -> Optional[str]:
        """
        Extracts the first email address found in the resume text.
        
        Email regex breakdown:
        [a-zA-Z0-9._%+-]+  : username part (letters, digits, dots, underscores, %, +, -)
        @                   : literal @ symbol
        [a-zA-Z0-9.-]+      : domain name
        \.                  : literal dot
        [a-zA-Z]{2,}        : TLD (com, org, io, etc.) — minimum 2 chars
        """
        # re.search finds first match anywhere in string
        # re.IGNORECASE makes matching case-insensitive
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(email_pattern, text, re.IGNORECASE)

        # Return matched string or None if no email found
        return match.group(0).lower() if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """
        Extracts phone number from resume text.
        
        Handles multiple common formats:
        - (555) 123-4567
        - 555-123-4567
        - 555.123.4567
        - +1 555 123 4567
        - 5551234567
        
        Regex breakdown:
        (\+\d{1,3}[\s-]?)? : optional country code (+1, +44, etc.)
        \(?\d{3}\)?        : area code, optional parentheses
        [\s.-]?            : optional separator (space, dot, dash)
        \d{3}              : first 3 digits
        [\s.-]?            : optional separator
        \d{4}              : last 4 digits
        """
        phone_pattern = r'(\+\d{1,3}[\s-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}'
        match = re.search(phone_pattern, text)
        return match.group(0).strip() if match else None

    def _extract_name(self, text: str) -> Optional[str]:
        """
        Attempts to extract candidate name from the top of the resume.
        
        Strategy: The first non-empty line of a resume is almost always
        the candidate's name. We validate it:
        - Must have 2+ words (first and last name)
        - Must contain only letters, spaces, hyphens, apostrophes
        - Must not be a section header (EXPERIENCE, EDUCATION, etc.)
        - Must be under 50 characters
        
        This is a heuristic — not 100% accurate. Module 3's NER will
        provide better name extraction using spaCy.
        """
        # Known section headers to exclude — if first line matches these,
        # it's not a name
        section_headers = {
            'resume', 'curriculum vitae', 'cv', 'objective', 'summary',
            'profile', 'experience', 'education', 'skills', 'contact'
        }

        # Split text into lines and examine first several
        lines = text.split('\n')

        for line in lines[:5]:  # Check first 5 lines only
            # Clean the candidate line
            candidate = line.strip()

            # Must have content
            if not candidate:
                continue

            # Must be reasonably short (names aren't paragraphs)
            if len(candidate) > 50:
                continue

            # Must contain only valid name characters
            # Allows: letters, spaces, hyphens (O'Brien), apostrophes (O'Brien)
            if not re.match(r'^[a-zA-Z\s\-\'\.]+$', candidate):
                continue

            # Must have at least 2 words (first + last name)
            words = candidate.split()
            if len(words) < 2:
                continue

            # Must not be a section header
            if candidate.lower() in section_headers:
                continue

            # Passed all checks — likely a name
            # title() capitalizes first letter of each word
            return candidate.title()

        # Could not extract name — NLP pipeline will try again
        return None


# Instantiate a single parser for use across the application
# Stateless design means one instance handles all requests safely
resume_parser = ResumeParser()